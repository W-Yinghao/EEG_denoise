# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

doc = Document("中文版本.docx")

def iter_block_items(parent):
    from docx.document import Document as _Doc
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    if isinstance(parent, _Doc):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Count overall stats
nparas = len(doc.paragraphs)
ntables = len(doc.tables)
print(f"### STATS: paragraphs={nparas}, tables={ntables}")
print("="*70)

# Check for OMML math
body_xml = doc.element.body.xml
n_omml = body_xml.count('oMath')
print(f"### OMML math element count (rough): {n_omml}")
print("="*70)

# Heading outline
print("### HEADING OUTLINE:")
for p in doc.paragraphs:
    st = p.style.name if p.style else ""
    if st and ("Heading" in st or "Title" in st or "标题" in st):
        print(f"  [{st}] {p.text}")
print("="*70)
